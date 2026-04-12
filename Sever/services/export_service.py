"""Export service: convert Markdown content to DOCX or PDF."""

from __future__ import annotations

import os
import re
import tempfile
from pathlib import Path


# ---------------------------------------------------------------------------
# Markdown → DOCX
# ---------------------------------------------------------------------------

def markdown_to_docx(md_text: str, output_path: str, md_base_dir: str | None = None) -> None:
    """Convert Markdown text to a styled DOCX file at *output_path*.

    Handles: headings (H1-H6), paragraphs, bold/italic, inline code, code
    blocks, unordered/ordered lists, horizontal rules, blockquotes, and images.

    md_base_dir: directory of the source Markdown file, used to resolve
    relative image paths for embedding via doc.add_picture().
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- document-level styles ------------------------------------------------
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    def _set_heading_style(paragraph, level: int):
        sizes = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 11))
        if level <= 2:
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0x75)

    def _apply_inline(paragraph, text: str):
        """Add inline-formatted runs to *paragraph* from a text fragment."""
        # Patterns: **bold**, *italic*, `code`, and plain text
        pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+?)`)')
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                paragraph.add_run(text[last:m.start()])
            full = m.group(0)
            if full.startswith('**'):
                r = paragraph.add_run(m.group(2))
                r.bold = True
            elif full.startswith('*'):
                r = paragraph.add_run(m.group(3))
                r.italic = True
            else:  # backtick code
                r = paragraph.add_run(m.group(4))
                r.font.name = "Courier New"
                r.font.size = Pt(10)
            last = m.end()
        if last < len(text):
            paragraph.add_run(text[last:])

    lines = md_text.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # ---- fenced code block ----
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(18)
                pPr = p._p.get_or_add_pPr()
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F5F5F5")
                pPr.append(shd)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ---- heading ----
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            p = doc.add_paragraph()
            p.add_run(text)
            _set_heading_style(p, level)
            i += 1
            continue

        # ---- horizontal rule ----
        if re.match(r'^(\-{3,}|\*{3,}|_{3,})\s*$', line):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # ---- blockquote ----
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            _apply_inline(p, text)
            p.paragraph_format.left_indent = Pt(18)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                run.italic = True
            i += 1
            continue

        # ---- unordered list ----
        ul_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if ul_match:
            text = ul_match.group(2).strip()
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, text)
            i += 1
            continue

        # ---- ordered list ----
        ol_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if ol_match:
            text = ol_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            _apply_inline(p, text)
            i += 1
            continue

        # ---- image ----
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_src = img_match.group(2).strip()
            inserted = False
            if md_base_dir and img_src:
                lower_src = img_src.lower()
                if not lower_src.startswith(("http://", "https://", "data:", "file://")):
                    abs_img = os.path.normpath(os.path.join(md_base_dir, img_src))
                    if os.path.isfile(abs_img):
                        try:
                            doc.add_picture(abs_img, width=Inches(6))
                            inserted = True
                        except Exception:
                            pass
            if not inserted and alt_text:
                p = doc.add_paragraph()
                r = p.add_run(f"[图片: {alt_text}]")
                r.italic = True
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # ---- blank line ----
        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        # ---- normal paragraph ----
        p = doc.add_paragraph()
        _apply_inline(p, line.strip())
        i += 1

    doc.save(output_path)


# ---------------------------------------------------------------------------
# Markdown → PDF  (via Playwright – persistent browser singleton)
# ---------------------------------------------------------------------------

# Absolute path to locally bundled KaTeX (from the frontend node_modules).
# Using local files removes the CDN round-trip that was a major source of delay.
_HERE = os.path.dirname(os.path.abspath(__file__))
_KATEX_DIR = os.path.normpath(
    os.path.join(_HERE, "..", "..", "View", "node_modules", "katex", "dist")
)
_KATEX_CSS = os.path.join(_KATEX_DIR, "katex.min.css")
_KATEX_JS  = os.path.join(_KATEX_DIR, "katex.min.js")
_KATEX_AR  = os.path.join(_KATEX_DIR, "contrib", "auto-render.min.js")

# Fall back to CDN if local files are not found (e.g. node_modules not installed)
_USE_LOCAL_KATEX = all(os.path.isfile(p) for p in (_KATEX_CSS, _KATEX_JS, _KATEX_AR))


def _build_katex_tags() -> str:
    """Return HTML <link>/<script> tags that load KaTeX."""
    if _USE_LOCAL_KATEX:
        css_href = Path(_KATEX_CSS).as_uri()
        js_src   = Path(_KATEX_JS).as_uri()
        ar_src   = Path(_KATEX_AR).as_uri()
    else:
        base = "https://cdn.jsdelivr.net/npm/katex@0.16.9/dist"
        css_href = f"{base}/katex.min.css"
        js_src   = f"{base}/katex.min.js"
        ar_src   = f"{base}/contrib/auto-render.min.js"

    return (
        f'<link rel="stylesheet" href="{css_href}">\n'
        f'<script src="{js_src}"></script>\n'
        f'<script src="{ar_src}"'
        ' onload="renderMathInElement(document.body,{delimiters:['
        '{left:\'$$\',right:\'$$\',display:true},'
        '{left:\'$\',right:\'$\',display:false}'
        ']})"></script>\n'
    )


_PDF_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;700&display=swap">
<style>
  body {{
    font-family: "Noto Sans SC", "Noto Sans CJK SC", "WenQuanYi Micro Hei", "Segoe UI", Arial, "Liberation Sans", sans-serif;
    font-size: 13px;
    line-height: 1.7;
    max-width: 820px;
    margin: 40px auto;
    padding: 0 24px;
    color: #222;
  }}
  h1,h2,h3,h4,h5,h6 {{ color: #1a5675; margin-top: 1.4em; }}
  h1 {{ font-size: 1.9em; }} h2 {{ font-size: 1.5em; }} h3 {{ font-size: 1.25em; }}
  code {{
    background: #f5f5f5; border-radius: 3px;
    padding: 2px 5px; font-family: "Courier New", monospace; font-size: 0.88em;
  }}
  pre {{
    background: #f5f5f5; border-radius: 5px;
    padding: 12px 16px; overflow-x: auto;
  }}
  pre code {{ background: none; padding: 0; }}
  blockquote {{
    border-left: 4px solid #ccc; margin: 0;
    padding-left: 16px; color: #555; font-style: italic;
  }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #ccc; padding: 6px 12px; }}
  th {{ background: #f0f0f0; font-weight: bold; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 1.5em 0; }}
  img {{ max-width: 100%; height: auto; display: block; margin: 0.5em 0; }}
</style>
{bilingual_style}
{katex_tags}
</head>
<body>
{body}
</body>
</html>
"""

# ---------------------------------------------------------------------------
# Bilingual-mode CSS — mirrors MarkdownViewer.vue bilingual scoped styles.
# Accepts hue/saturation/intensity so PDF colour matches the user's theme
# selection (forwarded from the browser's localStorage via query params).
# Default values (195 / 70 / 6) match the app's built-in default preset.
# ---------------------------------------------------------------------------
def _build_bilingual_css(hue: int = 195, saturation: int = 70, intensity: int = 6, font_size: int = 15) -> str:
    """Return a <style> block with bilingual colours and font size derived from user params.

    hue:        colour wheel position (0-360)
    saturation: HSL saturation percent (0-100)
    intensity:  background opacity level (2-15), maps to alpha the same way
                the frontend does: blockquote bg = intensity*0.01,
                [译] label bg = intensity*0.02
    font_size:  base font size in px (12-20); overrides the template body font-size
                so that PDF matches the on-screen reading size exactly.
    """
    bg_alpha = round(intensity * 0.01, 4)
    label_alpha = round(intensity * 0.02, 4)
    return f"""\
<style>
  /* Override body font-size to match user's on-screen reading preference */
  body {{
    font-size: {font_size}px !important;
  }}
  /* English source paragraph (blockquote) */
  blockquote {{
    border-left: 3px solid hsl({hue}, {saturation}%, 45%);
    background: hsla({hue}, {saturation}%, 50%, {bg_alpha});
    color: #555;
    font-size: 0.92em;
    font-style: normal;
    padding: 0.5em 0.85em;
    margin: 0.5em 0 0;
    border-radius: 0 6px 6px 0;
  }}
  /* [译] label — strong tag that is the only child of its paragraph */
  p > strong:only-child {{
    display: inline-block;
    font-size: 0.72em;
    font-weight: 600;
    letter-spacing: 0.04em;
    color: hsl({hue}, {saturation}%, 38%);
    background: hsla({hue}, {saturation}%, 50%, {label_alpha});
    border: 1px solid hsla({hue}, {saturation}%, 50%, 0.28);
    border-radius: 4px;
    padding: 0.1em 0.45em;
    margin: 0.25em 0 0.1em;
    line-height: 1.5;
  }}
  /* Chinese translation paragraph — inherits body font-size (= user's chosen value) */
  p {{
    line-height: 1.8;
    margin: 0.2em 0 0.6em;
  }}
  /* Heading sizes — same em ratios as the frontend bilingual-mode overrides
     so that PDF output matches the on-screen visual hierarchy exactly. */
  h1 {{ font-size: 1.61em; }}
  h2 {{ font-size: 1.33em; }}
  h3 {{ font-size: 1.17em; }}
  h4, h5, h6 {{ font-size: 1.06em; }}
  /* Section separator */
  hr {{
    border: none;
    border-top: 1px dashed #ccc;
    margin: 0.85em 0;
    opacity: 0.6;
  }}
</style>
"""


def _rewrite_html_img_srcs(html: str, base_dir: str) -> str:
    """Rewrite relative <img src="..."> paths in *html* to absolute file:// URIs.

    Only src values that are not already absolute paths or URLs are rewritten.
    Uses pathlib.Path.as_uri() for cross-platform compatibility (Windows/Linux).
    """
    def _fix_src(m: re.Match) -> str:
        quote = m.group(1)
        src = m.group(2)
        if not src:
            return m.group(0)
        lower = src.lower()
        if lower.startswith(("http://", "https://", "data:", "file://", "//")):
            return m.group(0)
        abs_path = os.path.normpath(os.path.join(base_dir, src))
        try:
            file_uri = Path(abs_path).as_uri()
        except ValueError:
            return m.group(0)
        return f"src={quote}{file_uri}{quote}"

    return re.sub(r'src=(["\'])([^"\']*)\1', _fix_src, html, flags=re.IGNORECASE)


def markdown_to_pdf(
    md_text: str,
    output_path: str,
    md_base_dir: str | None = None,
    bilingual: bool = False,
    bilingual_hue: int = 195,
    bilingual_saturation: int = 70,
    bilingual_intensity: int = 6,
    bilingual_font_size: int = 15,
) -> None:
    """Convert Markdown text to a PDF file at *output_path* via Playwright.

    A fresh Playwright Chromium instance is created per call to avoid the
    thread-binding issue with sync_playwright (Playwright's sync API binds its
    internal event loop to the thread that created it; FastAPI runs sync
    endpoints on a thread-pool where threads may differ between requests,
    causing "cannot switch to a different thread" errors with a shared
    singleton).  The cold-start overhead (~1-2 s) is acceptable for a
    user-triggered download action.

    KaTeX is loaded from local node_modules when available, removing the CDN
    round-trip delay.

    md_base_dir: directory of the source Markdown file, used to resolve
    relative image paths (e.g. mineru_bundle/images/xxx.jpg) to absolute
    file:// URIs so Playwright can load them. Pass os.path.dirname(md_path).

    bilingual: when True, injects bilingual-mode CSS that mirrors the
    MarkdownViewer.vue bilingual styles (coloured blockquote accent,
    [译] label badge, dashed hr separators).

    bilingual_hue / bilingual_saturation / bilingual_intensity: HSL colour
    parameters forwarded from the user's browser theme preference so that the
    downloaded PDF matches the on-screen colour.  Defaults match the app's
    built-in default preset (hue=195, saturation=70, intensity=6).

    bilingual_font_size: base font size in px (12-20) forwarded from the
    user's on-screen reading preference so that PDF text size matches
    exactly what the user sees in the browser.  Default is 15px.
    """
    import markdown as md_lib
    from playwright.sync_api import sync_playwright

    html_body = md_lib.markdown(
        md_text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )

    if md_base_dir:
        html_body = _rewrite_html_img_srcs(html_body, md_base_dir)

    full_html = _PDF_HTML_TEMPLATE.format(
        bilingual_style=_build_bilingual_css(bilingual_hue, bilingual_saturation, bilingual_intensity, bilingual_font_size) if bilingual else "",
        katex_tags=_build_katex_tags(),
        body=html_body,
    )

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".html", encoding="utf-8", delete=False
    ) as tmp_html:
        tmp_html.write(full_html)
        tmp_html_path = tmp_html.name

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            try:
                page = browser.new_page()
                try:
                    page.goto(Path(tmp_html_path).as_uri())
                    # Wait for network idle so that external resources finish loading:
                    # - Google Fonts (Noto Sans SC) needed for CJK characters in PDF
                    # - KaTeX CDN scripts (when local node_modules not present)
                    try:
                        page.wait_for_load_state("networkidle", timeout=10000)
                    except Exception:
                        # Timeout is acceptable — proceed with whatever loaded
                        pass
                    if _USE_LOCAL_KATEX:
                        try:
                            page.wait_for_function(
                                "() => !document.querySelector('.katex-error')"
                                " && document.readyState === 'complete'",
                                timeout=3000,
                            )
                        except Exception:
                            pass
                    page.pdf(
                        path=output_path,
                        format="A4",
                        margin={"top": "30mm", "bottom": "30mm", "left": "25mm", "right": "25mm"},
                        print_background=True,
                    )
                finally:
                    page.close()
            finally:
                browser.close()
    finally:
        try:
            os.unlink(tmp_html_path)
        except OSError:
            pass
